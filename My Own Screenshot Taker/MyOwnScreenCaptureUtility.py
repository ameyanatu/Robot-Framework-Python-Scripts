import pyscreenshot as ImageGrab
from docx import Document
from docx.shared import Inches

def take_cool_screenshot(testCaseNumber, content):
    document = Document()
    document.add_heading(testCaseNumber,level=2)
    document.add_paragraph(content)
    ImageGrab.grab_to_file(testCaseNumber + '.png')
    document.add_picture(testCaseNumber + '.png',width=Inches(5.90551), height=Inches(3.54331))
    document.save(testCaseNumber+'ScreenShot'+'.docx')

if __name__ == "__main__":
    take_cool_screenshot("12345","Test")